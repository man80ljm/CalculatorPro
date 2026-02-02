import os
import requests
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QIcon
from PyQt6.QtWidgets import (
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QTextEdit,
    QFrame,
    QScrollArea,
    QWidget,
    QVBoxLayout,
    QComboBox,
)
from utils import get_outputs_dir


class TestApiThread(QThread):
    result = pyqtSignal(str)

    def __init__(self, api_key: str, model: str):
        super().__init__()
        self.api_key = api_key
        self.model = model

    def run(self):
        self.result.emit(test_api(self.api_key, self.model))


def test_api(api_key: str, model: str) -> str:
    """测试 API 连接 - 支持 DeepSeek 和 Anthropic"""
    api_key = api_key.strip().strip("<").strip(">")
    
    if model.startswith("claude"):
        # Test Anthropic API
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=api_key)
            message = client.messages.create(
                model=model,
                max_tokens=10,
                messages=[{"role": "user", "content": "测试连接"}]
            )
            return "连接成功"
        except ImportError:
            return "连接失败: 未安装 anthropic 库，请运行 pip install anthropic"
        except Exception as exc:
            return f"连接失败: {exc}"
    else:
        # Test DeepSeek API
        url = "https://api.deepseek.com/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": "测试连接"},
            ],
            "temperature": 0.7,
            "top_p": 1,
            "max_tokens": 10,
            "stream": False,
        }
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=10)
            response.raise_for_status()
            return "连接成功"
        except requests.RequestException as exc:
            return f"连接失败: {exc}"


def test_deepseek_api(api_key: str) -> str:
    """向后兼容的 DeepSeek 测试函数"""
    return test_api(api_key, "deepseek-chat")


class SettingsDialog(QDialog):
    """设置弹窗：API Key、课程简介、目标要求、上一学年达成度"""

    def __init__(
        self,
        parent=None,
        api_key="",
        ai_model="deepseek-chat",
        description="",
        objective_requirements=None,
        objectives_count=0,
        previous_achievement_file="",
    ):
        super().__init__(parent)
        self.setWindowTitle("设置")
        self.resize(650, 700)
        self.setMinimumSize(650, 500)
        self.setWindowIcon(QIcon(os.path.join(os.path.dirname(__file__), "..", "calculator.ico")))

        self.api_key_value = api_key
        self.ai_model_value = ai_model
        self.description_value = description
        self.objective_requirements = objective_requirements or []
        self.objectives_count = objectives_count
        self.previous_achievement_file = previous_achievement_file
        self.objective_inputs = []
        self._build_ui()
        # allow vertical resize

    def _build_ui(self):
        outer_layout = QVBoxLayout()
        outer_layout.setSpacing(0)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        layout = QVBoxLayout()
        layout.setSpacing(10)
        layout.setContentsMargins(16, 16, 16, 16)
        self.setStyleSheet("""
            QDialog { 
                background: #FFFFFF;
            }
            QLabel {
                font-size: 14px;
                color: #333333;
                font-weight: 600;
            }
            QLineEdit {
                background: #FFFFFF;
                border: 2px solid #C0C0C0;
                border-radius: 8px;
                padding: 8px 10px;
                font-size: 13px;
                color: #333333;
            }
            QLineEdit:focus {
                border: 2px solid #8A8A8A;
            }
            QTextEdit {
                background: #FFFFFF;
                border: 2px solid #C0C0C0;
                border-radius: 8px;
                padding: 8px 10px;
                font-size: 13px;
                color: #333333;
            }
            QTextEdit::viewport {
                border: 2px solid #C0C0C0;
                border-radius: 8px;
                background: #FFFFFF;
            }
            QTextEdit:focus {
                border: 2px solid #8A8A8A;
            }
            QPushButton {
                background: #8A8A8A;
                color: #FFFFFF;
                border: none;
                border-radius: 10px;
                padding: 10px 18px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #7C7C7C;
            }
            QPushButton:pressed {
                background: #6E6E6E;
            }
        """)

        title = QLabel("设置")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("font-size: 18px; font-weight: bold;")
        layout.addWidget(title)

        desc_label = QLabel("课程简介:")
        self.desc_input = QTextEdit()
        self.desc_input.setFixedHeight(140)
        self.desc_input.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.desc_input.setFrameStyle(QFrame.Shape.Box | QFrame.Shadow.Plain)
        self.desc_input.setLineWidth(2)
        self.desc_input.setText(self.description_value)
        self.desc_input.setPlaceholderText("请输入课程简介...")
        self.desc_input.setStyleSheet("border: 2px solid #C0C0C0; border-radius: 8px;")
        layout.addWidget(desc_label)
        layout.addWidget(self.desc_input)

        self._build_objective_inputs(layout)

        self.import_prev_btn = QPushButton("导入上一学年达成度表")
        self.import_prev_btn.clicked.connect(self._import_previous_file)
        layout.addWidget(self.import_prev_btn)

        self.file_path_label = QLabel(self.previous_achievement_file or "未选择文件")
        layout.addWidget(self.file_path_label)

        # AI Model selection
        model_layout = QHBoxLayout()
        model_label = QLabel("AI 模型:")
        self.model_combo = QComboBox()
        self.model_combo.addItems([
            "deepseek-chat",
            "claude-3-5-sonnet-20241022",
            "claude-3-5-sonnet-20240620",
            "claude-3-opus-20240229",
            "claude-3-sonnet-20240229",
            "claude-3-haiku-20240307",
        ])
        # Set current model
        index = self.model_combo.findText(self.ai_model_value)
        if index >= 0:
            self.model_combo.setCurrentIndex(index)
        model_layout.addWidget(model_label)
        model_layout.addWidget(self.model_combo)
        layout.addLayout(model_layout)

        api_layout = QHBoxLayout()
        api_label = QLabel("API KEY:")
        self.api_input = QLineEdit()
        self.api_input.setText(self.api_key_value)
        self.api_input.setPlaceholderText("请输入 API Key (DeepSeek 或 Anthropic)")
        self.test_btn = QPushButton("检测")
        self.test_btn.clicked.connect(self._test_api)
        api_layout.addWidget(api_label)
        api_layout.addWidget(self.api_input)
        api_layout.addWidget(self.test_btn)
        layout.addLayout(api_layout)

        btn_layout = QHBoxLayout()
        save_btn = QPushButton("保存")
        clear_btn = QPushButton("清空")
        save_btn.setFixedWidth(120)
        clear_btn.setFixedWidth(120)
        save_btn.clicked.connect(self._on_save)
        clear_btn.clicked.connect(self._on_clear)
        btn_layout.addStretch()
        btn_layout.addWidget(save_btn)
        btn_layout.addSpacing(214)
        btn_layout.addWidget(clear_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        content = QWidget()
        content.setLayout(layout)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setWidget(content)
        outer_layout.addWidget(scroll)
        self.setLayout(outer_layout)

    def _build_objective_inputs(self, layout):
        if self.objectives_count <= 0:
            return
        section = QLabel("课程目标要求:")
        section.setStyleSheet("font-weight: bold;")
        layout.addWidget(section)
        for idx in range(self.objectives_count):
            label = QLabel(f"课程目标{idx + 1}要求")
            input_box = QLineEdit()
            input_box.setPlaceholderText(f"请输入目标{idx + 1}要求")
            if idx < len(self.objective_requirements):
                input_box.setText(self.objective_requirements[idx])
            layout.addWidget(label)
            layout.addWidget(input_box)
            self.objective_inputs.append(input_box)

    def _import_previous_file(self):
        file_name, _ = QFileDialog.getOpenFileName(
            self,
            "选择上一学年达成度表",
            "",
            "Excel Files (*.xlsx)",
        )
        if file_name:
            self.previous_achievement_file = file_name
            self.file_path_label.setText(file_name)

    def _test_api(self):
        api_key = self.api_input.text().strip()
        if not api_key:
            QMessageBox.warning(self, "提示", "请先输入API Key")
            return
        model = self.model_combo.currentText()
        self.test_dialog = QMessageBox(self)
        self.test_dialog.setWindowTitle("测试连接")
        self.test_dialog.setText("连接中...")
        self.test_dialog.setStandardButtons(QMessageBox.StandardButton.NoButton)
        self.test_dialog.show()
        self.test_thread = TestApiThread(api_key, model)
        self.test_thread.result.connect(self._on_test_result)
        self.test_thread.start()

    def _on_test_result(self, result: str):
        self.test_dialog.setText(result)
        self.test_dialog.setStandardButtons(QMessageBox.StandardButton.Ok)
        self.test_dialog.exec()


    def _set_cell_border(self, cell, size=4):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tcPr.append(borders)
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_tag = qn(f'w:{edge}')
            element = borders.find(edge_tag)
            if element is None:
                element = OxmlElement(f'w:{edge}')
                borders.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), str(size))  # 0.5pt -> 4
            element.set(qn('w:color'), '000000')

    def _export_course_basic_word(self, data: dict):
        root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        output_dir = get_outputs_dir()
        os.makedirs(output_dir, exist_ok=True)
        course_name = (data.get('course_name') or '\u8bfe\u7a0b').strip()
        safe_name = re.sub(r'[\\/:*?"<>|]', '_', course_name)
        output_path = os.path.join(output_dir, "1.\u8bfe\u7a0b\u57fa\u672c\u4fe1\u606f\u8868.docx")

        doc = Document()
        table = doc.add_table(rows=4, cols=6)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        col_width = Cm(14.64 / 6)

        rows = [
            ('\u8bfe\u7a0b\u540d\u79f0', data.get('course_name', ''), '\u5b66\u5206', data.get('credits', ''), '\u5b66\u65f6', data.get('hours', '')),
            ('\u8bfe\u7a0b\u6027\u8d28', data.get('course_type', ''), '\u8bfe\u7a0b\u4ee3\u7801', data.get('course_code', ''), '\u5b66\u5e74\u5b66\u671f', data.get('school_year_term', '')),
            ('\u5f00\u8bfe\u5b66\u9662', data.get('college', ''), '\u4efb\u8bfe\u6559\u5e08', data.get('teacher', ''), '\u4e0a\u8bfe\u4e13\u4e1a', data.get('major', '')),
            ('\u4e0a\u8bfe\u73ed\u7ea7', data.get('class_name', ''), '\u4e0a\u8bfe\u4eba\u6570', data.get('student_count', ''), '\u8003\u6838\u4eba\u6570', data.get('exam_count', '')),
        ]

        for r_idx, row_vals in enumerate(rows):
            row = table.rows[r_idx]
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            for c_idx in range(6):
                cell = row.cells[c_idx]
                cell.width = col_width
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(row_vals[c_idx]) if c_idx < len(row_vals) else '')
                run.font.name = '\u4eff\u5b8b'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '\u4eff\u5b8b')
                run.font.size = Pt(12)
                if c_idx in (0, 2, 4):
                    run.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_cell_border(cell, size=4)

        doc.save(output_path)
        return output_path



    def _export_grad_req_docx(self, grad_req_map):
        root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        output_dir = get_outputs_dir()
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(
            output_dir,
            "3.\u8bfe\u7a0b\u76ee\u6807\u4e0e\u6bd5\u4e1a\u8981\u6c42\u7684\u5bf9\u5e94\u5173\u7cfb\u8868.docx",
        )

        doc = Document()
        rows_count = max(1, len(grad_req_map)) + 1
        table = doc.add_table(rows=rows_count, cols=3)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        total_cm = 14.64
        col1 = 2.79
        col2 = 3.37
        col3 = total_cm - col1 - col2
        col_widths = [col1, col2, col3]

        headers = [
            "\u8bfe\u7a0b\u76ee\u6807",
            "\u652f\u6491\u7684\u6bd5\u4e1a\u8981\u6c42",
            "\u652f\u6491\u7684\u6bd5\u4e1a\u8981\u6c42\u6307\u6807\u70b9",
        ]
        for c_idx, title in enumerate(headers):
            cell = table.cell(0, c_idx)
            cell.width = Cm(col_widths[c_idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(title)
            run.font.name = "\u4eff\u5b8b"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "\u4eff\u5b8b")
            run.font.size = Pt(12)
            run.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            self._set_cell_border(cell, size=4)

        for r in range(1, rows_count):
            obj_name = f"\u8bfe\u7a0b\u76ee\u6807{r}"
            requirement = ""
            indicator = ""
            if r - 1 < len(grad_req_map):
                row = grad_req_map[r - 1]
                obj_name = row.get('objective', obj_name)
                requirement = row.get('requirement', '')
                indicator = row.get('indicator', '')

            for c_idx, val in enumerate([obj_name, requirement, indicator]):
                cell = table.cell(r, c_idx)
                cell.width = Cm(col_widths[c_idx])
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(str(val))
                run.font.name = "\u4eff\u5b8b"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "\u4eff\u5b8b")
                run.font.size = Pt(12)
                run.bold = (c_idx == 0)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_cell_border(cell, size=4)
                table.rows[r].height = Cm(1)
                table.rows[r].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        table.rows[0].height = Cm(1)
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        doc.save(output_path)
        return output_path

    def _on_save(self):
        try:
            new_api = self.api_input.text().strip()
            if new_api:
                self.api_key_value = new_api
            self.ai_model_value = self.model_combo.currentText()
            self.description_value = self.desc_input.toPlainText().strip()
            self.objective_requirements = [item.text().strip() for item in self.objective_inputs]

            parent = self.parent()
            export_error = None
            if parent:
                parent.api_key = self.api_key_value
                parent.ai_model = self.ai_model_value
                parent.course_description = self.description_value
                parent.objective_requirements = self.objective_requirements
                parent.previous_achievement_file = self.previous_achievement_file
                if hasattr(parent, "save_config"):
                    parent.save_config()

                try:
                    self._export_course_basic_word(getattr(parent, 'course_basic_info', {}) or {})
                except Exception as exc:
                    export_error = str(exc)

                try:
                    self._export_grad_req_docx(getattr(parent, 'grad_req_map', []) or [])
                except Exception as exc:
                    export_error = (export_error + ' ; ' if export_error else '') + str(exc)

            if export_error:
                QMessageBox.warning(self, "提示", f"保存成功，但生成课程基本信息表失败: {export_error}")
            else:
                QMessageBox.information(self, "提示", "保存成功")
            self.accept()
        except Exception as exc:
            QMessageBox.warning(self, "提示", f"保存失败：{exc}")

    def _on_clear(self):
        try:
            self.desc_input.clear()
            for item in self.objective_inputs:
                item.clear()
            self.file_path_label.setText("未选择文件")
            self.previous_achievement_file = ""

            parent = self.parent()
            if parent:
                if hasattr(parent, "usual_ratio"):
                    parent.usual_ratio = 0.0
                    parent.midterm_ratio = 0.0
                    parent.final_ratio = 0.0
                if hasattr(parent, "relation_payload"):
                    parent.relation_payload = None
                if hasattr(parent, "num_objectives"):
                    parent.num_objectives = 0
                if hasattr(parent, "course_open_info"):
                    parent.course_open_info = {}
                if hasattr(parent, "course_basic_info"):
                    parent.course_basic_info = {}
                if hasattr(parent, "grad_req_map"):
                    parent.grad_req_map = []
                parent.course_description = ""
                parent.objective_requirements = []
                parent.previous_achievement_file = ""
                if hasattr(parent, "save_config"):
                    parent.save_config()

            QMessageBox.information(self, "提示", "已清空")
        except Exception as exc:
            QMessageBox.warning(self, "提示", f"清空失败：{exc}")

