"""
ReportBuilder - 用于将多个Word文档拼接到模板中
"""
import os
import glob
from docx import Document
from copy import deepcopy
from datetime import datetime
from utils import get_outputs_dir


class ReportBuilder:
    """报告构建器 - 将多个Word文档内容合并到模板中"""
    
    def __init__(self, template_path, output_dir):
        """
        初始化报告构建器
        
        Args:
            template_path: 模板文件路径
            output_dir: 输出目录
        """
        self.template_path = template_path
        self.output_dir = output_dir
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    def build(self, course_open_info, course_basic_info, achievement_data):
        """
        构建完整报告
        
        Args:
            course_open_info: 开课信息字典
            course_basic_info: 课程基本信息字典
            achievement_data: 达成度数据字典
            
        Returns:
            str: 输出文件路径
        """
        # 加载模板
        doc = Document(self.template_path)
        
        # 1. 替换基本文本占位符
        self._replace_text_placeholders(doc, course_open_info, course_basic_info)
        
        # 2. 插入各个文档的内容
        self._insert_documents(doc)
        
        # 3. 保存输出文件
        output_path = self._generate_output_path(course_open_info)
        doc.save(output_path)
        
        return output_path
    
    def _replace_text_placeholders(self, doc, course_open_info, course_basic_info):
        """替换文本占位符"""
        # 获取当前日期
        now = datetime.now()
        
        # 构建替换映射
        replacements = {
            '{{year_start}}': course_open_info.get('year_start', str(now.year)),
            '{{year_end}}': course_open_info.get('year_end', str(now.year + 1)),
            '{{term}}': course_open_info.get('term', '1'),
            '{{course_name}}': course_basic_info.get('course_name', ''),
            '{{department}}': course_open_info.get('department', ''),
            '{{teacher}}': course_open_info.get('teacher', ''),
            '{{report_date}}': f"{now.year} 年 {now.month} 月"
        }
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    inline = para.runs
                    for run in inline:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                for run in para.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
    
    def _find_doc_by_prefix(self, outputs_dir, prefix_pattern):
        """
        根据模式查找文档（智能匹配下划线和点号）
        
        Args:
            outputs_dir: 输出目录
            prefix_pattern: 文件模式（如 "1_课程基本信息表.docx"）
        
        Returns:
            str: 找到的文件路径，未找到返回None
        """
        # 1. 直接精确匹配
        direct_path = os.path.join(outputs_dir, prefix_pattern)
        if os.path.exists(direct_path):
            print(f"    [精确匹配] {prefix_pattern}")
            return direct_path
        
        # 2. 提取数字前缀（如 "1"）和主要内容（如 "课程基本信息表"）
        # 支持格式: "1_xxx.docx" 或 "1.xxx.docx" 或 "1 xxx.docx"
        import re
        match = re.match(r'^(\d+)[._\s]*(.+?)\.docx$', prefix_pattern)
        if not match:
            return None
        
        number = match.group(1)  # 提取数字
        main_text = match.group(2)  # 提取主要文本
        
        # 3. 尝试多种分隔符组合
        separators = ['_', '.', ' ', '-']
        for sep in separators:
            # 尝试: 数字 + 分隔符 + 文本 + .docx
            variant = f"{number}{sep}{main_text}.docx"
            variant_path = os.path.join(outputs_dir, variant)
            if os.path.exists(variant_path):
                print(f"    [变体匹配] {variant}")
                return variant_path
        
        # 4. 使用通配符模糊匹配（匹配以数字开头的文件）
        patterns = [
            f"{number}_*{main_text}*.docx",
            f"{number}.*{main_text}*.docx",
            f"{number}*{main_text}*.docx",
        ]
        
        for pattern in patterns:
            full_pattern = os.path.join(outputs_dir, pattern)
            matches = glob.glob(full_pattern)
            if matches:
                print(f"    [模糊匹配] 模式: {pattern}, 找到: {os.path.basename(matches[0])}")
                return matches[0]
        
        # 5. 最后的尝试：只匹配数字开头
        fallback_pattern = os.path.join(outputs_dir, f"{number}*.docx")
        matches = glob.glob(fallback_pattern)
        if matches:
            # 过滤掉明显不相关的文件（如时间戳）
            for match in matches:
                basename = os.path.basename(match)
                # 排除包含时间戳的文件（如 20260130_223114）
                if not re.search(r'\d{8}_\d{6}', basename):
                    print(f"    [回退匹配] {basename}")
                    return match
        
        return None
    
    def _insert_documents(self, doc):
        """在模板中插入各个文档的内容"""
        # 获取outputs目录（统一到应用根目录）
        outputs_dir = get_outputs_dir()
        
        print(f"\n{'='*60}")
        print(f"[ReportBuilder] 查找文档目录: {outputs_dir}")
        print(f"[ReportBuilder] 目录存在: {os.path.exists(outputs_dir)}")
        
        if os.path.exists(outputs_dir):
            files = os.listdir(outputs_dir)
            print(f"[ReportBuilder] 找到 {len(files)} 个文件:")
            for f in files:
                if f.endswith('.docx'):
                    print(f"  - {f}")
        print(f"{'='*60}\n")
        
        # 定义要插入的文档（使用下划线格式，匹配函数会智能处理）
        doc_mappings = {
            '{{INSERT_DOC_1}}': '1_课程基本信息表.docx',
            '{{INSERT_DOC_2}}': '2_课程成绩统计表.docx',
            '{{INSERT_DOC_3}}': '3_课程目标与毕业要求的对应关系表.docx',
            '{{INSERT_DOC_4}}': '4_课程考核与课程目标对应关系表.docx',
            '{{INSERT_DOC_5}}': '5_基于考核结果的课程目标达成情况评价结果表.docx',
            '{{INSERT_DOC_6}}': '6_课程目标达成情况分析_存在问题及改进措施表.docx'
        }
        
        # 收集需要处理的段落（从后往前处理，避免索引混乱）
        paragraphs_to_process = []
        for i, para in enumerate(doc.paragraphs):
            for placeholder, prefix in doc_mappings.items():
                if placeholder in para.text:
                    print(f"[占位符 {placeholder}]")
                    print(f"  查找模式: {prefix}")
                    # 使用智能匹配查找文档
                    doc_path = self._find_doc_by_prefix(outputs_dir, prefix)
                    filename = os.path.basename(doc_path) if doc_path else f"{prefix}"
                    paragraphs_to_process.append((i, para, placeholder, doc_path, filename))
                    
                    if doc_path:
                        print(f"  ✓ 找到文件: {filename}")
                    else:
                        print(f"  ✗ 未找到文件")
                    print()
        
        # 从后往前处理（避免插入导致索引变化）
        for i, para, placeholder, doc_path, filename in reversed(paragraphs_to_process):
            if doc_path and os.path.exists(doc_path):
                print(f"→ 插入文档: {filename}")
                # 清空占位符段落
                para.clear()
                # 插入文档内容
                self._insert_document_content(doc, para, doc_path)
            else:
                print(f"→ 警告: 跳过缺失文档 {filename}")
                para.text = f"[缺失文档: {filename}]"
    
    def _insert_document_content(self, target_doc, insert_point_para, source_doc_path):
        """
        在指定位置插入源文档的内容
        
        使用深拷贝方式，直接复制XML元素
        """
        try:
            # 加载源文档
            source_doc = Document(source_doc_path)
            
            # 获取插入点
            main_body = target_doc.element.body
            insert_after_element = insert_point_para._element
            insert_index = main_body.index(insert_after_element)
            
            # 插入所有元素（段落和表格）
            offset = 1
            for element in source_doc.element.body:
                # 深拷贝元素
                new_element = deepcopy(element)
                # 插入到目标位置
                main_body.insert(insert_index + offset, new_element)
                offset += 1
            
            print(f"  ✓ 成功插入 {offset-1} 个元素")
            
            # 添加一个空段落作为分隔
            separator = target_doc.add_paragraph()._element
            main_body.insert(insert_index + offset, separator)
            
        except Exception as e:
            print(f"  ✗ 插入失败: {e}")
            import traceback
            traceback.print_exc()
            # 在出错位置插入错误提示
            insert_point_para.text = f"[文档插入失败: {os.path.basename(source_doc_path)} - {str(e)}]"
    
    def _generate_output_path(self, course_open_info):
        """生成输出文件路径"""
        course_name = course_open_info.get('course_name', '课程')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{course_name}_达成度报告_{timestamp}.docx"
        return os.path.join(self.output_dir, filename)


# 测试代码
if __name__ == '__main__':
    print("ReportBuilder 模块已加载")
