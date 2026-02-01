import os
import numpy as np
import pandas as pd
import json
import re
import requests
from typing import List, Dict,Callable, Optional
from openpyxl.styles import Alignment, PatternFill, Font, Border, Side
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
from apply_noise import GradeReverseEngine
from utils import normalize_score, get_grade_level, calculate_final_score, calculate_achievement_level, adjust_column_widths, get_outputs_dir
import time
import random
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

class AIReportMixin:
        def _set_cell_border(self, cell, size=4, color='000000'):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            for edge in ('top', 'left', 'bottom', 'right'):
                elem = tcBorders.find(qn(f'w:{edge}'))
                if elem is None:
                    elem = OxmlElement(f'w:{edge}')
                    tcBorders.append(elem)
                if size:
                    elem.set(qn('w:val'), 'single')
                    elem.set(qn('w:sz'), str(size))
                    elem.set(qn('w:space'), '0')
                    elem.set(qn('w:color'), color)
                else:
                    elem.set(qn('w:val'), 'nil')


        def test_deepseek_api(self, api_key: str) -> str:
            """测试 DeepSeek API 连接"""
            url = "https://api.deepseek.com/v1/chat/completions"
            api_key = api_key.strip().strip('<').strip('>')
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            payload = {
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": "测试连接"}
                ],
                "temperature": 0.7,
                "top_p": 1,
                "max_tokens": 10,
                "stream": False
            }
            
            try:
                response = requests.post(url, headers=headers, json=payload, timeout=10)
                response.raise_for_status()
                return "连接成功"
            except requests.RequestException as e:
                error_message = f"连接失败: {str(e)}"
                if hasattr(e, 'response') and e.response is not None:
                    error_message += f"\n服务器返回: {e.response.text}"
                return error_message

        def call_deepseek_api(self, prompt: str) -> str:
            """调用 DeepSeek API 获取答案，包含重试与超时处理。"""
            if not self.api_key:
                return "请先设置API Key"

            url = "https://api.deepseek.com/v1/chat/completions"
            api_key = self.api_key.strip().strip("<").strip(">")
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            }

            max_tokens = 600
            match = re.search(r"接近(\d+)字", prompt)
            if match:
                try:
                    word_limit = int(match.group(1))
                    max_tokens = max(200, min(1500, word_limit * 4))
                except ValueError:
                    max_tokens = 600

            payload = {
                "model": "deepseek-chat",
                "messages": [
                    {"role": "system", "content": "You are a helpful assistant specializing in course analysis and improvement."},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.7,
                "top_p": 1,
                "max_tokens": max_tokens,
                "stream": False,
            }

            max_retries = 3
            for attempt in range(max_retries):
                try:
                    response = requests.post(url, headers=headers, json=payload, timeout=30)
                    response.raise_for_status()
                    return response.json()["choices"][0]["message"]["content"].strip()
                except requests.Timeout:
                    if attempt < max_retries - 1:
                        print(f"API 调用超时，正在重试（第 {attempt + 1}/{max_retries} 次）...")
                        time.sleep(2)
                        continue
                    return "API 调用超时，请检查网络连接或稍后重试（可能需要使用VPN或代理访问 api.deepseek.com）"
                except requests.RequestException as e:
                    error_message = f"API 调用失败: {str(e)}"
                    if hasattr(e, "response") and e.response is not None:
                        error_message += f"\n服务端返回: {e.response.text}"
                    if attempt < max_retries - 1:
                        print(f"API 调用失败，正在重试（第 {attempt + 1}/{max_retries} 次）...")
                        time.sleep(2)
                        continue
                    return error_message
                except (KeyError, IndexError):
                    return "API 返回格式错误，无法解析结果"

        def load_previous_achievement(self, file_path: str) -> None:
            """\u52a0\u8f7d\u4e0a\u4e00\u5b66\u5e74\u8fbe\u6210\u5ea6\u8868\uff0c\u63d0\u53d6\u5404\u8bfe\u7a0b\u76ee\u6807\u7684\u5206\u76ee\u6807\u8fbe\u6210\u503c\u4ee5\u53ca\u8bfe\u7a0b\u76ee\u6807\u8fbe\u6210\u503c\u3002"""
            def _objective_count():
                payload = self.relation_payload or {}
                objectives = payload.get("objectives") if isinstance(payload, dict) else None
                if objectives:
                    return len(objectives)
                if self.objective_requirements:
                    return len(self.objective_requirements)
                return 5

            def _init_defaults(n):
                data = {f'\u8bfe\u7a0b\u76ee\u6807{i}': 0 for i in range(1, n + 1)}
                data['\u8bfe\u7a0b\u603b\u76ee\u6807'] = 0
                return data

            obj_count = _objective_count()
            if not file_path:
                self.previous_achievement_data = _init_defaults(obj_count)
                return

            try:
                if not os.path.exists(file_path):
                    self.previous_achievement_data = _init_defaults(obj_count)
                    if self.status_label:
                        self.status_label.setText("\u672a\u627e\u5230\u4e0a\u4e00\u5b66\u5e74\u8fbe\u6210\u5ea6\u8868\uff0c\u5df2\u4f7f\u7528\u9ed8\u8ba4\u503c")
                    return

                xls = pd.ExcelFile(file_path)
                df = None
                for sheet in xls.sheet_names:
                    tmp = pd.read_excel(file_path, sheet_name=sheet)
                    cols = [str(c).strip() for c in tmp.columns]
                    if "\u8bfe\u7a0b\u5206\u76ee\u6807" in cols or "\u8bfe\u7a0b\u76ee\u6807" in cols:
                        df = tmp
                        break
                if df is None:
                    df = pd.read_excel(file_path)

                cols = [str(c).strip() for c in df.columns]
                data = _init_defaults(obj_count)

                if "\u8bfe\u7a0b\u5206\u76ee\u6807" in cols and "\u5206\u76ee\u6807\u8fbe\u6210\u503c" in cols:
                    for i in range(1, obj_count + 1):
                        key = f'\u8bfe\u7a0b\u76ee\u6807{i}'
                        rows = df[df["\u8bfe\u7a0b\u5206\u76ee\u6807"].astype(str).str.strip() == key]
                        if not rows.empty:
                            val = rows["\u5206\u76ee\u6807\u8fbe\u6210\u503c"].dropna().tolist()
                            if val:
                                data[key] = float(val[0])
                    total_rows = df[df["\u8bfe\u7a0b\u5206\u76ee\u6807"].astype(str).str.strip().isin([
                        "\u8bfe\u7a0b\u76ee\u6807\u8fbe\u6210\u503c",
                        "\u8bfe\u7a0b\u603b\u76ee\u6807\u8fbe\u6210\u503c",
                        "\u8bfe\u7a0b\u603b\u8fbe\u6210\u503c",
                    ])]
                    if not total_rows.empty:
                        val = total_rows["\u5206\u76ee\u6807\u8fbe\u6210\u503c"].dropna().tolist()
                        if val:
                            data["\u8bfe\u7a0b\u603b\u76ee\u6807"] = float(val[0])
                elif "\u8bfe\u7a0b\u76ee\u6807" in cols:
                    value_col = None
                    for cand in [
                        "\u4e0a\u4e00\u5e74\u5ea6\u8fbe\u6210\u5ea6",
                        "\u4e0a\u4e00\u8f6e\u6559\u5b66\u5206\u76ee\u6807\u8fbe\u6210\u503c",
                        "\u5206\u76ee\u6807\u8fbe\u6210\u503c",
                    ]:
                        if cand in cols:
                            value_col = cand
                            break
                    if value_col:
                        for _, row in df.iterrows():
                            target = str(row["\u8bfe\u7a0b\u76ee\u6807"]).strip()
                            if target in data and isinstance(row[value_col], (int, float)):
                                data[target] = float(row[value_col])
                        total_rows = df[df["\u8bfe\u7a0b\u76ee\u6807"].astype(str).str.strip().isin([
                            "\u8bfe\u7a0b\u76ee\u6807\u8fbe\u6210\u503c",
                            "\u8bfe\u7a0b\u603b\u76ee\u6807\u8fbe\u6210\u503c",
                            "\u8bfe\u7a0b\u603b\u8fbe\u6210\u503c",
                        ])]
                        if not total_rows.empty and isinstance(total_rows[value_col].iloc[0], (int, float)):
                            data["\u8bfe\u7a0b\u603b\u76ee\u6807"] = float(total_rows[value_col].iloc[0])
                self.previous_achievement_data = data
                if self.status_label:
                    self.status_label.setText(f"\u5df2\u52a0\u8f7d\u4e0a\u4e00\u5b66\u5e74\u8fbe\u6210\u5ea6\u8868: {os.path.basename(file_path)}")
            except Exception as e:
                if self.status_label:
                    self.status_label.setText("\u52a0\u8f7d\u4e0a\u4e00\u5b66\u5e74\u8fbe\u6210\u5ea6\u8868\u5931\u8d25")
                raise ValueError(f"\u52a0\u8f7d\u4e0a\u4e00\u5b66\u5e74\u8fbe\u6210\u5ea6\u8868\u5931\u8d25: {str(e)}")

        def generate_improvement_report(self, answers: list[str] | None, output_dir: str | None = None) -> str:
            base_dir = Path(output_dir) if output_dir else Path(get_outputs_dir())
            base_dir.mkdir(parents=True, exist_ok=True)

            course_name = getattr(self, 'course_name', '')
            if not course_name and hasattr(self, 'course_name_input'):
                try:
                    course_name = self.course_name_input.text().strip()
                except Exception:
                    course_name = ''
            course_name = course_name or '课程'
            safe_name = re.sub(r'[\/:*?"<>|]', '_', course_name)
            output_file = base_dir / f"6.课程目标达成情况分析、存在问题及改进措施表.docx"

            # 删除旧版 xlsx 兼容文件
            old_xlsx = base_dir / f"{safe_name}课程分目标达成情况分析、存在问题及改进措施.xlsx"
            if old_xlsx.exists():
                old_xlsx.unlink(missing_ok=True)

            obj_count = len(self.objective_requirements or [])
            total_questions = 1 + obj_count * 2
            answers = answers or []
            while len(answers) < total_questions:
                answers.append('')

            overall_answer = answers[0].strip() if answers else ''

            rows: list[tuple[str, str]] = []
            rows.append(('（一）总体情况', 'heading'))
            rows.append((overall_answer, 'answer'))
            rows.append(('（二）课程分目标达成情况分析、存在问题及改进措施', 'heading'))

            idx = 1
            for i in range(1, obj_count + 1):
                rows.append((f"{i}. 课程目标{i}", 'fixed'))
                rows.append(('（1）达成情况分析：', 'fixed'))
                rows.append((answers[idx].strip() if idx < len(answers) else '', 'answer'))
                idx += 1
                rows.append(('（2）存在问题及改进措施：', 'fixed'))
                rows.append((answers[idx].strip() if idx < len(answers) else '', 'answer'))
                idx += 1

            doc = Document()
            table = doc.add_table(rows=len(rows), cols=2)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            total_width_cm = 14.64
            left_col_cm = 1.0
            right_col_cm = total_width_cm - left_col_cm
            table.columns[0].width = Cm(left_col_cm)
            table.columns[1].width = Cm(right_col_cm)

            fixed_size = Pt(15)  # 标题字体
            answer_size = Pt(14)  # 正文字体

            for r_idx, (text, kind) in enumerate(rows):
                row = table.rows[r_idx]
                row.cells[0].text = ''
                p = row.cells[1].paragraphs[0]
                p.text = ''
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.first_line_indent = None
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)

                run = p.add_run(text)
                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
                if kind == 'answer':
                    run.font.size = answer_size
                    run.bold = False
                elif kind == 'heading':
                    run.font.size = fixed_size
                    run.bold = True
                else:
                    run.font.size = fixed_size
                    run.bold = False

                row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
                row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

            for row in table.rows:
                for cell in row.cells:
                    self._set_cell_border(cell, size=0)

            doc.save(output_file)
            return str(output_file)

        def store_api_key(self, api_key: str) -> None:
            """存储 API Key"""
            self.api_key = api_key
            if hasattr(self, 'status_label') and self.status_label:
                self.status_label.setText("已存储API Key")

        def generate_ai_report(self, *args, **kwargs) -> None:
            """适配旧接口调用的 generate_improvement_report 包装器"""
            answers = kwargs.get('answers') if isinstance(kwargs, dict) else None
            self.generate_improvement_report(answers=answers)
