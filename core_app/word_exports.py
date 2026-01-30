import os
import numpy as np
import pandas as pd
import json
import re
import requests
from typing import List, Dict,Callable, Optional
from openpyxl.styles import Alignment, PatternFill, Font, Border, Side
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
from apply_noise import GradeReverseEngine
from utils import normalize_score, get_grade_level, calculate_final_score, calculate_achievement_level, adjust_column_widths
import time
import random
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH

class WordExportMixin:
        def _set_docx_cell_margins(self, cell, top=0, bottom=0, left=0, right=0):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                tcMar = OxmlElement('w:tcMar')
                tcPr.append(tcMar)
            for tag, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
                node = tcMar.find(qn(f'w:{tag}'))
                if node is None:
                    node = OxmlElement(f'w:{tag}')
                    tcMar.append(node)
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')

        def _set_docx_cell_border(self, cell, size=4):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.find(qn('w:tcBorders'))
            if borders is None:
                borders = OxmlElement('w:tcBorders')
                tcPr.append(borders)
            for edge in ('top', 'left', 'bottom', 'right'):
                edge_tag = qn(f'w:{edge}')
                element = borders.find(edge_tag)
                if element is None:
                    element = OxmlElement(f'w:{edge}')
                    borders.append(element)
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), str(size))  # 0.5pt -> 4
                element.set(qn('w:color'), '000000')

        def _export_stats_docx(self, composition_text, max_score, min_score, avg_score, counts, ratios):
            root = os.path.abspath(os.path.dirname(__file__))
            output_dir = os.path.join(root, 'outputs')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, '\u0032.\u8bfe\u7a0b\u6210\u7ee9\u7edf\u8ba1\u8868.docx')

            doc = Document()
            table = doc.add_table(rows=5, cols=6)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            total_cm = 14.64
            first_cm = 3.75
            other_cm = (total_cm - first_cm) / 5

            data = [
                # 第一行：成绩构成（假设第一列是标题）
                ["成绩构成", composition_text.strip(), "", "", "", ""],
                
                # 第二行：最高成绩
                ["最高成绩", max_score, "最低成绩", min_score, "平均成绩", avg_score],
                
                # 第三行：成绩等级（这里是报错的重点，必须用 \n 换行）
                ["成绩等级", "90-100\n(优秀)", "80-89\n(良好)", "70-79\n(中等)", "60-69\n(及格)", "<60\n(不及格)"],
                
                # 第四行：人数
                ["人数"] + list(counts),
                
                # 第五行：占考核人数的比例
                ["占考核人数的比例"] + [f"{r*100:.2f}%" for r in ratios],
            ]

            bold_coords = {
                (0,0),
                (1,0), (1,2), (1,4),
                (2,0), (2,1), (2,2), (2,3), (2,4), (2,5),
                (3,0),
                (4,0),
            }

            # === 循环部分修改 ===
            for r_idx, row_vals in enumerate(data):
                row = table.rows[r_idx]
                for c_idx in range(6):
                    # 【核心修复】：对于第1行(r_idx=0)，只处理前两列，跳过后面被合并的列
                    if r_idx == 0 and c_idx > 1:
                        continue

                    cell = row.cells[c_idx]
                    cell.width = Cm(first_cm if c_idx == 0 else other_cm)
                    cell.text = "" # 清除可能存在的默认标记
                    
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    
                    # 获取要写入的内容
                    text_content = str(row_vals[c_idx]) if c_idx < len(row_vals) else ""
                    run = p.add_run(text_content)
                    
                    # 设置字体
                    run.font.name = "\u4eff\u5b8b"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "\u4eff\u5b8b") # 这里的中文名对应 FangSong
                    run.font.size = Pt(12)
                    
                    if (r_idx, c_idx) in bold_coords:
                        run.bold = True
                    
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    self._set_docx_cell_border(cell, size=4)
                    self._set_docx_cell_margins(cell, top=0, bottom=0, left=0, right=0)

            # 最后执行合并，这时候被合并的单元格是干净的，不会有空行
            table.cell(0,1).merge(table.cell(0,5))
            
            doc.save(output_path)
            return output_path

        def _export_eval_result_docx(self, links, obj_keys, method_avgs, prev_data,
                                     total_attainment, expected_attainment, prev_total):
            root = os.path.abspath(os.path.dirname(__file__))
            output_dir = os.path.join(root, 'outputs')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(
                output_dir,
                "5.基于考核结果的课程目标达成情况评价结果表.docx"
            )

            doc = Document()
            table = doc.add_table(rows=1, cols=7)
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            total_cm = 14.64
            col_w = Cm(total_cm / 7)

            headers = [
                "课程分目标",
                "考核环节",
                "分权重",
                "分值/满分",
                "学生实际得分平均分",
                "分目标达成值",
                "上一轮教学分目标达成值",
            ]

            header_row = table.rows[0]
            for c_idx, text in enumerate(headers):
                cell = header_row.cells[c_idx]
                cell.width = col_w
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(text)
                run.font.name = "\u4eff\u5b8b"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "\u4eff\u5b8b")
                run.font.size = Pt(12)
                run.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_docx_cell_border(cell, size=4)
                self._set_docx_cell_margins(cell, top=0, bottom=0, left=0, right=0)

            trPr = header_row._tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "1")
            trPr.append(tblHeader)

            row_map = []
            for idx, obj_key in enumerate(obj_keys):
                obj_name = f"课程目标{idx + 1}"
                obj_start = len(table.rows)

                obj_weight_sum = 0.0
                obj_actual_sum = 0.0

                for link in links:
                    link_name = link.get("name", "")
                    if "平时" in link_name:
                        display_link = "平时成绩"
                    elif "期中" in link_name:
                        display_link = "期中考核"
                    elif "期末" in link_name:
                        display_link = "期末考核"
                    else:
                        display_link = link_name

                    link_ratio = float(link.get("ratio", 0))
                    methods = link.get("methods", []) or []

                    support_sum = 0.0
                    actual_sum = 0.0
                    for m in methods:
                        supports = m.get("supports", {}) or {}
                        weight = float(supports.get(obj_key, 0))
                        support_sum += weight
                        m_name = m.get("name")
                        m_avg = float(method_avgs.get(m_name, 0))
                        actual_sum += m_avg * weight

                    target_weight = link_ratio * 100.0 * support_sum
                    actual_score = link_ratio * actual_sum

                    obj_weight_sum += target_weight
                    obj_actual_sum += actual_score

                    row = table.add_row()
                    values = [
                        obj_name if len(table.rows) - 1 == obj_start else "",
                        display_link,
                        f"{round(target_weight, 2)}",
                        "100",
                        f"{round(actual_score, 2)}",
                        "",
                        "",
                    ]

                    for c_idx, val in enumerate(values):
                        cell = row.cells[c_idx]
                        cell.width = col_w
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        run = p.add_run(str(val))
                        run.font.name = "\u4eff\u5b8b"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "\u4eff\u5b8b")
                        run.font.size = Pt(12)
                        if c_idx in (0, 1):
                            run.bold = True
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        self._set_docx_cell_border(cell, size=4)
                        self._set_docx_cell_margins(cell, top=0, bottom=0, left=0, right=0)

                obj_row_count = len(table.rows) - obj_start
                row_map.append((obj_start, obj_row_count, obj_weight_sum, obj_actual_sum, obj_name))

            # ... (前面的代码不变)

            for obj_start, obj_row_count, obj_weight_sum, obj_actual_sum, obj_name in row_map:
                if obj_row_count <= 0:
                    continue
                
                # 1. 计算达成值
                achievement = round(obj_actual_sum / obj_weight_sum, 3) if obj_weight_sum > 0 else 0
                prev_val = prev_data.get(obj_name, 0) if prev_data else 0
                prev_val = 0 if prev_val is None else prev_val

                # 2. 如果需要合并（即该目标跨越多行）
                if obj_row_count > 1:
                    # === 关键修改 ===
                    # 合并 第0列(课程目标名称), 第5列(达成值), 第6列(上一轮)
                    for col_idx in [0, 5, 6]:
                        start_cell = table.cell(obj_start, col_idx)
                        end_cell = table.cell(obj_start + obj_row_count - 1, col_idx)
                        
                        # 执行合并
                        start_cell.merge(end_cell)
                        
                        # === 清理合并后产生的多余段落 ===
                        # merge后，start_cell.text 可能会变成 "目标1\n\n\n" 这种形式
                        # 我们要强制清空，只写我们需要的一个值
                        
                        text_to_write = ""
                        if col_idx == 0:
                            text_to_write = obj_name
                        elif col_idx == 5:
                            text_to_write = str(achievement)
                        elif col_idx == 6:
                            text_to_write = str(prev_val)
                        
                        # 1. 清除所有内容
                        start_cell.text = "" 
                        
                        # 2. 添加唯一的段落和Run
                        p = start_cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        
                        run = p.add_run(text_to_write)
                        run.font.name = "FangSong"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                        run.font.size = Pt(12)
                        run.bold = (col_idx == 0) # 只有第1列加粗
                        
                        start_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                else:
                    # 如果只有一行（不需要合并），直接填值
                    first_row = table.rows[obj_start]
                    
                    # 填达成值
                    cell_ach = first_row.cells[5]
                    cell_ach.text = "" # 先清空防止追加
                    p = cell_ach.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(str(achievement))
                    run.font.name = "FangSong"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                    run.font.size = Pt(12)
                    cell_ach.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # 填上一轮值
                    cell_prev = first_row.cells[6]
                    cell_prev.text = "" # 先清空
                    p = cell_prev.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(str(prev_val))
                    run.font.name = "FangSong"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                    run.font.size = Pt(12)
                    cell_prev.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # ... (前面的 obj_start 循环代码保持不变) ...

            # === 以下是底部汇总行的修复代码 ===
            summary_rows = [
                ("课程目标达成值", total_attainment),
                ("课程目标达成期望值", expected_attainment),
                ("上一轮教学课程目标达成值", prev_total),
            ]
            
            for label, value in summary_rows:
                row = table.add_row()
                r_idx = len(table.rows) - 1
                
                # 1. 先给所有格子刷一遍基础样式（边框、宽度），防止合并后边框丢失
                for c_idx in range(7):
                    cell = row.cells[c_idx]
                    cell.width = col_w
                    self._set_docx_cell_border(cell, size=4)
                    self._set_docx_cell_margins(cell, top=0, bottom=0, left=0, right=0)
                
                # 2. 处理左边的大格子 (第0-4列合并)
                cell_label = table.cell(r_idx, 0)
                cell_label.merge(table.cell(r_idx, 4))
                
                # 【关键】清空合并后产生的多余段落
                cell_label.text = "" 
                
                # 重新写入标签
                p = cell_label.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(label)
                run.font.name = "FangSong"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                run.font.size = Pt(12)
                run.bold = True
                cell_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # 3. 处理右边的数值格子 (第5-6列合并)
                cell_value = table.cell(r_idx, 5)
                cell_value.merge(table.cell(r_idx, 6))
                
                # 【关键】清空合并后产生的多余段落
                cell_value.text = ""
                
                # 重新写入数值
                p = cell_value.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                val_str = str(value if value is not None else 0)
                run = p.add_run(val_str)
                run.font.name = "FangSong"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                run.font.size = Pt(12)
                run.bold = False
                cell_value.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            doc.save(output_path)
            return output_path

        def _export_grad_req_docx(self, grad_req_map):
            from docx import Document
            from docx.shared import Cm, Pt
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            root = os.path.abspath(os.path.dirname(__file__))
            output_dir = os.path.join(root, 'outputs')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, '3.课程目标与毕业要求的对应关系表.docx')

            doc = Document()
            
            # 定义列宽（厘米）
            col_widths_cm = [2.79, 3.37, 8.48]
            total_cm = sum(col_widths_cm)
            
            # 计算百分比
            col_percentages = [int((w / total_cm) * 5000) for w in col_widths_cm]
            
            # 创建表格
            rows_count = max(1, len(grad_req_map)) + 1
            table = doc.add_table(rows=rows_count, cols=3)
            tbl = table._element
            
            # 设置表格属性
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # 清除旧设置
            for child in list(tblPr):
                if child.tag in (qn('w:tblW'), qn('w:tblLayout')):
                    tblPr.remove(child)
            
            # 设置表格宽度为 100%
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            tblPr.append(tblW)
            
            # 设置固定布局
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            
            # 设置表格网格（这个保留，定义列的比例）
            tblGrid = tbl.find(qn('w:tblGrid'))
            if tblGrid is None:
                tblGrid = OxmlElement('w:tblGrid')
                tbl.insert(1, tblGrid)
            else:
                for child in list(tblGrid):
                    tblGrid.remove(child)
            
            # 添加网格列
            for pct in col_percentages:
                gridCol = OxmlElement('w:gridCol')
                gridCol.set(qn('w:w'), str(pct))
                tblGrid.append(gridCol)
            
            # ===【关键修复】删除所有单元格的 tcW 宽度设置===
            # 这样就相当于取消了"指定宽度"的勾选
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    
                    # 删除所有 tcW 元素（这就是取消"指定宽度"勾选）
                    for tcW in list(tcPr.findall(qn('w:tcW'))):
                        tcPr.remove(tcW)
                    
                    # 如果 tcPr 为空，也可以删除它
                    if len(tcPr) == 0:
                        tc.remove(tcPr)
            
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            headers = ["课程目标", "支撑的毕业要求", "支撑的毕业要求指标点"]
            for c_idx, title in enumerate(headers):
                cell = table.cell(0, c_idx)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(title)
                run.font.name = "\u4eff\u5b8b"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                run.font.size = Pt(12)
                run.bold = True
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._set_docx_cell_border(cell, size=4)

            # 数据行
            rows_data = grad_req_map if grad_req_map else []
            for r_idx in range(1, len(table.rows)):
                obj_name = f"课程目标{r_idx}"
                requirement = ""
                indicator = ""
                
                if r_idx - 1 < len(rows_data):
                    row = rows_data[r_idx - 1]
                    obj_name = row.get('objective', obj_name)
                    requirement = row.get('requirement', "")
                    indicator = row.get('indicator', "")

                for c_idx, val in enumerate([obj_name, requirement, indicator]):
                    cell = table.cell(r_idx, c_idx)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(str(val))
                    run.font.name = "\u4eff\u5b8b"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "仿宋")
                    run.font.size = Pt(12)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    self._set_docx_cell_border(cell, size=4)
                
                table.rows[r_idx].height = Cm(1)
                table.rows[r_idx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            table.rows[0].height = Cm(1)
            table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            doc.save(output_path)
            return output_path
