import json
import os
from typing import List, Tuple, Optional

from PyQt6.QtCore import Qt
from PyQt6.QtGui import QIntValidator, QKeySequence, QIcon
from PyQt6.QtWidgets import (
    QApplication,
    QAbstractItemView,
    QDialog,
    QGridLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QTableWidget,
    QTableWidgetItem,
    QVBoxLayout,
    QStyledItemDelegate,
    QMenu,
)

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None


def _ensure_outputs_dir(base_dir: str) -> str:
    outputs_dir = os.path.join(base_dir, "outputs")
    os.makedirs(outputs_dir, exist_ok=True)
    return outputs_dir


def _format_percent(value: float) -> str:
    text = f"{value:.2f}".rstrip("0").rstrip(".")
    return f"{text}%"


def _parse_percent_text(text: str) -> Tuple[Optional[float], Optional[str]]:
    if text is None:
        return None, "输入不能为空"
    cleaned = text.strip().replace("％", "%")
    if not cleaned:
        return None, "输入不能为空"
    if not cleaned.endswith("%"):
        return None, "请输入带%符号的百分数"
    number_part = cleaned[:-1].strip()
    try:
        value = float(number_part)
    except ValueError:
        return None, "百分数格式不正确"
    if value < 0 or value > 100:
        return None, "百分数必须在0-100之间"
    return value, None


def _set_cell_readonly(item: QTableWidgetItem) -> None:
    item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)


def _set_cell_bg(item: QTableWidgetItem, color: str) -> None:
    item.setBackground(Qt.GlobalColor.transparent)
    item.setBackground(Qt.GlobalColor.lightGray if color == "gray" else Qt.GlobalColor.red)


class PasteTableWidget(QTableWidget):
    def __init__(self, *args, paste_callback=None, **kwargs):
        super().__init__(*args, **kwargs)
        self._paste_callback = paste_callback
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.DefaultContextMenu)

    def keyPressEvent(self, event):
        if event.matches(QKeySequence.StandardKey.Paste):
            self.paste_from_clipboard()
            return
        super().keyPressEvent(event)

    def paste_from_clipboard(self):
        text = QApplication.clipboard().text()
        if not text:
            return
        start_row, start_col = self._get_paste_anchor()
        if start_row is None:
            return
        rows = [row for row in text.splitlines() if row.strip() != ""]
        if not rows:
            return
        self.blockSignals(True)
        try:
            for r_offset, row_text in enumerate(rows):
                cols = row_text.split("\t")
                for c_offset, value in enumerate(cols):
                    r = start_row + r_offset
                    c = start_col + c_offset
                    if r >= self.rowCount() or c >= self.columnCount():
                        continue
                    item = self.item(r, c)
                    if item is None:
                        item = QTableWidgetItem()
                        self.setItem(r, c, item)
                    if not (item.flags() & Qt.ItemFlag.ItemIsEditable):
                        continue
                    item.setText(value.strip())
        finally:
            self.blockSignals(False)
        if self._paste_callback:
            self._paste_callback()

    def _get_paste_anchor(self):
        ranges = self.selectedRanges()
        if ranges:
            top = min(r.topRow() for r in ranges)
            left = min(r.leftColumn() for r in ranges)
            return top, left
        start = self.currentIndex()
        if not start.isValid():
            return None, None
        return start.row(), start.column()

    def contextMenuEvent(self, event):
        menu = QMenu(self)
        paste_action = menu.addAction("粘贴")
        action = menu.exec(event.globalPos())
        if action == paste_action:
            self.paste_from_clipboard()


class PercentItemDelegate(QStyledItemDelegate):
    def createEditor(self, parent, option, index):
        editor = QLineEdit(parent)
        editor.setValidator(QIntValidator(0, 100, editor))
        return editor

    def setEditorData(self, editor, index):
        text = index.data() or ""
        text = text.strip().replace("％", "%")
        if text.endswith("%"):
            text = text[:-1].strip()
        editor.setText(text)
        editor.selectAll()

    def setModelData(self, editor, model, index):
        text = editor.text().strip()
        if text == "":
            model.setData(index, "")
            return
        try:
            value = float(text)
        except ValueError:
            model.setData(index, "")
            return
        value = max(0.0, min(100.0, value))
        display = _format_percent(value)
        model.setData(index, display)


class RelationTableSetupDialog(QDialog):
    def __init__(self, parent=None, default_objectives=0, default_counts=None):
        super().__init__(parent)
        self.setWindowTitle("对应关系表")
        self.resize(300, 350)
        self.setWindowIcon(QIcon(os.path.join(os.path.dirname(__file__), "calculator.ico")))
        self.setStyleSheet("""
            QDialog {
                background: #FFFFFF;
                border: none;
                border-radius: 0px;
            }
            QLabel {
                font-size: 14px;
                color: #2C3E50;
                font-weight: 500;
            }
            QLineEdit {
                background: #F8F9FB;
                border: none;
                border-radius: 0px;
                padding: 10px 12px;
                font-size: 13px;
                color: #2C3E50;
            }
            QLineEdit:focus {
                border: 1px solid #BDBDBD;
                background: #FFFFFF;
            }
            QPushButton {
                background: #E0E0E0;
                border: none;
                border-radius: 10px;
                padding: 12px 32px;
                color: #333333;
                font-size: 14px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #D0D0D0;
            }
        """)
        self.result_values = None
        self._build_ui(default_objectives, default_counts)

    def _build_ui(self, default_objectives: int, default_counts):
        layout = QVBoxLayout()
        layout.setSpacing(20)
        layout.setContentsMargins(32, 32, 32, 32)

        # 标题
        title = QLabel("设置对应关系")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: #2C3E50;
            padding-bottom: 8px;
        """)
        layout.addWidget(title)

        grid = QGridLayout()
        grid.setHorizontalSpacing(12)
        grid.setVerticalSpacing(16)

        labels = ["课程目标:", "平时考核:", "期中考核:", "期末考核:"]
        self.inputs = []
        for row, label_text in enumerate(labels):
            label = QLabel(label_text)
            label.setFixedWidth(100)
            edit = QLineEdit()
            edit.setFixedWidth(120)
            edit.setValidator(QIntValidator(0, 20, self))
            edit.setPlaceholderText("0-20")
            if row == 0 and default_objectives:
                edit.setText(str(default_objectives))
            if default_counts and row > 0 and row - 1 < len(default_counts):
                edit.setText(str(default_counts[row - 1]))
            grid.addWidget(label, row, 0)
            grid.addWidget(edit, row, 1)
            self.inputs.append(edit)

        layout.addLayout(grid)

        btn = QPushButton("下一步")
        btn.clicked.connect(self._on_next)
        btn.setFixedHeight(44)
        layout.addWidget(btn, alignment=Qt.AlignmentFlag.AlignCenter)

        self.setLayout(layout)

    def _on_next(self):
        values = []
        for edit in self.inputs:
            text = edit.text().strip()
            if text == "":
                QMessageBox.warning(self, "提示", "请填写所有数量")
                return
            values.append(int(text))
        if values[0] <= 0:
            QMessageBox.warning(self, "提示", "课程目标数量必须大于0")
            return
        self.result_values = tuple(values)
        self.accept()


class RelationTableEditorDialog(QDialog):
    def __init__(self, parent, objectives_count: int, usual_count: int, midterm_count: int, final_count: int, existing_payload=None):
        super().__init__(parent)
        self.setWindowTitle("课程考核与课程目标对应关系表")
        self.setWindowIcon(QIcon(os.path.join(os.path.dirname(__file__), "calculator.ico")))
        self.setStyleSheet("""
            QDialog {
                background: #F5F7FA;
            }
            QTableWidget {
                background: #FFFFFF;
                border: none;
                border-radius: 0px;
                gridline-color: #E8ECF1;
            }
            QTableWidget::item {
                padding: 8px;
                color: #2C3E50;
            }
            QTableWidget::item:selected {
                background: #E3F2FD;
                color: #2C3E50;
            }
            QHeaderView::section {
                background: #F8F9FB;
                padding: 10px;
                border: 1px solid #E8ECF1;
                font-weight: 600;
                color: #2C3E50;
            }
            QPushButton {
                background: #6C757D;
                border: none;
                border-radius: 10px;
                padding: 10px 24px;
                color: white;
                font-weight: 500;
                min-width: 100px;
            }
            QPushButton:hover {
                background: #5A6268;
            }
        """)
        self.objectives_count = objectives_count
        self.link_counts = [usual_count, midterm_count, final_count]
        self.link_names = ["平时考核", "期中考核", "期末考核"]
        self.link_ratios = self._get_link_ratios()
        self.row_meta = []
        self._updating = False
        self.existing_payload = existing_payload

        if self.link_ratios is None:
            self.close()
            return

        self._build_ui()
        self._populate_table()
        self._apply_existing_payload()
        self._recalculate()
        self._resize_to_table()

    def _get_link_ratios(self) -> Optional[List[float]]:
        parent = self.parent()
        try:
            if hasattr(parent, "usual_ratio"):
                usual = float(parent.usual_ratio)
                midterm = float(parent.midterm_ratio)
                final = float(parent.final_ratio)
            else:
                usual = float(parent.usual_ratio_input.text().strip())
                midterm = float(parent.midterm_ratio_input.text().strip())
                final = float(parent.final_ratio_input.text().strip())
        except Exception:
            QMessageBox.warning(self, "提示", "请先填写平时/期中/期末占比(小数形式)")
            return None
        if abs((usual + midterm + final) - 1.0) != 0:
            QMessageBox.warning(self, "提示", "平时/期中/期末占比之和必须等于1")
            return None
        return [usual, midterm, final]

    def _build_ui(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)
        
        self.table = PasteTableWidget(paste_callback=self._recalculate)
        self.table.setEditTriggers(
            QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.EditKeyPressed
        )
        self.table.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectItems)
        self.table.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        self.table.cellChanged.connect(self._on_cell_changed)
        layout.addWidget(self.table)

        btn_layout = QHBoxLayout()
        back_btn = QPushButton("← 上一步")
        save_btn = QPushButton("💾 保存")
        
        back_btn.setStyleSheet("""
            QPushButton {
                background: #6C757D;
            }
            QPushButton:hover {
                background: #5A6268;
            }
        """)
        save_btn.setStyleSheet("""
            QPushButton {
                background: #28A745;
            }
            QPushButton:hover {
                background: #218838;
            }
        """)
        
        back_btn.clicked.connect(self._on_back)
        save_btn.clicked.connect(self._on_save)
        btn_layout.addWidget(back_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(save_btn)
        layout.addLayout(btn_layout)
        self.setLayout(layout)

    def keyPressEvent(self, event):
        if event.matches(QKeySequence.StandardKey.Paste):
            self.table.paste_from_clipboard()
            return
        super().keyPressEvent(event)

    def _populate_table(self):
        method_rows = []
        for link_idx, count in enumerate(self.link_counts):
            if count <= 0:
                method_rows.append({"link_idx": link_idx, "method_name": "无", "placeholder": True})
            else:
                for i in range(count):
                    method_rows.append({"link_idx": link_idx, "method_name": "", "placeholder": False})
        self.row_meta = method_rows

        cols = 2 + self.objectives_count + 2
        rows = 2 + len(method_rows) + 1
        self.table.setRowCount(rows)
        self.table.setColumnCount(cols)

        headers = ["考核环节", "考核方式"]
        headers += [f"课程目标{i+1}" for i in range(self.objectives_count)]
        headers += ["小计", "合计"]
        self.table.setHorizontalHeaderLabels(headers)

        # Header row 0 and 1
        self._set_item(0, 0, "考核环节", readonly=True, gray=True)
        self._set_item(0, 1, "考核方式", readonly=True, gray=True)
        self._set_item(0, 2, "课程目标分权重", readonly=True, gray=True)
        self.table.setSpan(0, 2, 1, self.objectives_count)
        self._set_item(0, 2 + self.objectives_count, "小计", readonly=True, gray=True)
        self._set_item(0, 3 + self.objectives_count, "合计", readonly=True, gray=True)

        # Merge header columns vertically
        for col in [0, 1, 2 + self.objectives_count, 3 + self.objectives_count]:
            self.table.setSpan(0, col, 2, 1)

        for i in range(self.objectives_count):
            self._set_item(1, 2 + i, f"课程目标{i+1}", readonly=True, gray=True)

        percent_delegate = PercentItemDelegate(self.table)
        for col in range(2, 2 + self.objectives_count):
            self.table.setItemDelegateForColumn(col, percent_delegate)

        # Data rows
        start_row = 2
        current_row = start_row
        for link_idx, link_name in enumerate(self.link_names):
            rows_for_link = [r for r in self.row_meta if r["link_idx"] == link_idx]
            first_row = current_row
            for row_info in rows_for_link:
                self._set_item(current_row, 1, row_info["method_name"], readonly=False, gray=False)
                if row_info["placeholder"]:
                    item = self.table.item(current_row, 1)
                    _set_cell_readonly(item)
                for obj_col in range(self.objectives_count):
                    self._set_item(current_row, 2 + obj_col, "0%", readonly=row_info["placeholder"], gray=False)
                self._set_item(current_row, 2 + self.objectives_count, "", readonly=True, gray=True)
                self._set_item(current_row, 3 + self.objectives_count, "", readonly=True, gray=True)
                current_row += 1

            link_label = self._format_link_label(link_name, self.link_ratios[link_idx])
            self._set_item(first_row, 0, link_label, readonly=True, gray=True)
            if len(rows_for_link) > 1:
                self.table.setSpan(first_row, 0, len(rows_for_link), 1)
                self.table.setSpan(first_row, 3 + self.objectives_count, len(rows_for_link), 1)

        # Total row
        total_row = rows - 1
        self._set_item(total_row, 0, "100%", readonly=True, gray=True)
        self._set_item(total_row, 1, "课程目标总权重", readonly=True, gray=True)
        for obj_col in range(self.objectives_count):
            self._set_item(total_row, 2 + obj_col, "", readonly=True, gray=True)
        self._set_item(total_row, 2 + self.objectives_count, "100%", readonly=True, gray=True)
        self._set_item(total_row, 3 + self.objectives_count, "100%", readonly=True, gray=True)

        self.table.resizeColumnsToContents()
        self.table.resizeRowsToContents()

    def _apply_existing_payload(self):
        payload = self.existing_payload or {}
        if not payload:
            return
        if payload.get("objectives_count") != self.objectives_count:
            return
        links = payload.get("links", [])
        if len(links) != len(self.link_names):
            return

        method_iters = [list(link.get("methods", [])) for link in links]
        method_pos = [0 for _ in self.link_names]

        for row_idx, meta in self._iter_method_rows():
            if meta.get("placeholder"):
                continue
            link_idx = meta["link_idx"]
            methods = method_iters[link_idx] if link_idx < len(method_iters) else []
            if method_pos[link_idx] >= len(methods):
                continue
            method = methods[method_pos[link_idx]]
            method_pos[link_idx] += 1

            name = (method.get("name") or "").strip()
            if name:
                item = self.table.item(row_idx, 1)
                if item is None:
                    item = QTableWidgetItem()
                    self.table.setItem(row_idx, 1, item)
                item.setText(name)

            supports = method.get("supports", {}) or {}
            for obj_idx in range(self.objectives_count):
                key = f"课程目标{obj_idx + 1}"
                value = supports.get(key)
                if value is None:
                    # 兼容意外 key
                    for k, v in supports.items():
                        if str(obj_idx + 1) in str(k):
                            value = v
                            break
                if value is None:
                    continue
                item = self.table.item(row_idx, 2 + obj_idx)
                if item is None:
                    item = QTableWidgetItem()
                    self.table.setItem(row_idx, 2 + obj_idx, item)
                item.setText(_format_percent(float(value) * 100))

            subtotal = method.get("subtotal")
            if subtotal is not None:
                subtotal_item = self.table.item(row_idx, 2 + self.objectives_count)
                if subtotal_item is None:
                    subtotal_item = QTableWidgetItem()
                    self.table.setItem(row_idx, 2 + self.objectives_count, subtotal_item)
                subtotal_item.setText(_format_percent(float(subtotal) * 100))

    def _resize_to_table(self):
        self.table.resizeColumnsToContents()
        self.table.resizeRowsToContents()
        width = self.table.verticalHeader().width() + self.table.horizontalHeader().length() + 60
        height = self.table.horizontalHeader().height()
        for row in range(self.table.rowCount()):
            height += self.table.rowHeight(row)
        height += 120  # 按钮和边距
        self.resize(width, height)
        self.setMinimumSize(width, height)

    def _set_item(self, row: int, col: int, text: str, readonly: bool, gray: bool):
        item = QTableWidgetItem(text)
        item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        if readonly:
            _set_cell_readonly(item)
        if gray:
            _set_cell_bg(item, "gray")
        self.table.setItem(row, col, item)

    def _format_link_label(self, name: str, ratio: float) -> str:
        percent = _format_percent(ratio * 100)
        return f"{name}\n{percent}"

    def _on_cell_changed(self, row: int, col: int):
        if self._updating:
            return
        self._recalculate()

    def _iter_method_rows(self):
        start_row = 2
        end_row = self.table.rowCount() - 1
        for idx in range(start_row, end_row):
            yield idx, self.row_meta[idx - start_row]

    def _recalculate(self):
        if self._updating:
            return
        self._updating = True
        self.table.blockSignals(True)
        try:
            obj_link_sums = [[0.0 for _ in range(self.objectives_count)] for _ in self.link_names]
            link_row_sums = [0.0 for _ in self.link_names]

            for row_idx, meta in self._iter_method_rows():
                row_sum = 0.0
                for obj_idx in range(self.objectives_count):
                    item = self.table.item(row_idx, 2 + obj_idx)
                    if item is None:
                        item = QTableWidgetItem("")
                        self.table.setItem(row_idx, 2 + obj_idx, item)
                    value, err = _parse_percent_text(item.text())
                    if err:
                        item.setBackground(Qt.GlobalColor.red)
                        continue
                    item.setBackground(Qt.GlobalColor.white)
                    row_sum += value
                    obj_link_sums[meta["link_idx"]][obj_idx] += value

                row_sum_item = self.table.item(row_idx, 2 + self.objectives_count)
                if row_sum_item is None:
                    row_sum_item = QTableWidgetItem("")
                    self.table.setItem(row_idx, 2 + self.objectives_count, row_sum_item)
                row_sum_item.setText(_format_percent(row_sum))
                row_sum_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                _set_cell_readonly(row_sum_item)
                row_sum_item.setBackground(Qt.GlobalColor.lightGray)
                link_row_sums[meta["link_idx"]] += row_sum

            # Set link totals in "合计" column
            start_row = 2
            for link_idx, count in enumerate(self.link_counts):
                rows_for_link = 1 if count <= 0 else count
                link_total_item = self.table.item(start_row, 3 + self.objectives_count)
                if link_total_item is None:
                    link_total_item = QTableWidgetItem("")
                    self.table.setItem(start_row, 3 + self.objectives_count, link_total_item)
                link_total_item.setText(_format_percent(link_row_sums[link_idx]))
                link_total_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                _set_cell_readonly(link_total_item)
                link_total_item.setBackground(Qt.GlobalColor.lightGray)
                start_row += rows_for_link

            # Compute objective total weights
            total_row = self.table.rowCount() - 1
            total_sum = 0.0
            for obj_idx in range(self.objectives_count):
                total = 0.0
                for link_idx, ratio in enumerate(self.link_ratios):
                    total += obj_link_sums[link_idx][obj_idx] * ratio
                total_sum += total
                item = self.table.item(total_row, 2 + obj_idx)
                if item is None:
                    item = QTableWidgetItem("")
                    self.table.setItem(total_row, 2 + obj_idx, item)
                item.setText(_format_percent(total))
                item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                _set_cell_readonly(item)
                item.setBackground(Qt.GlobalColor.lightGray)

            total_sum_item = self.table.item(total_row, 2 + self.objectives_count)
            if total_sum_item:
                total_sum_item.setText(_format_percent(total_sum))
        finally:
            self.table.blockSignals(False)
            self._updating = False

    def _validate(self) -> Optional[str]:
        # Check objective cell inputs
        for row_idx, meta in self._iter_method_rows():
            for obj_idx in range(self.objectives_count):
                item = self.table.item(row_idx, 2 + obj_idx)
                value, err = _parse_percent_text(item.text() if item else "")
                if err:
                    return f"第{row_idx-1}行,第{obj_idx+1}个课程目标:{err}"
                if value is None:
                    return f"第{row_idx-1}行,第{obj_idx+1}个课程目标:输入不能为空"

            row_sum_item = self.table.item(row_idx, 2 + self.objectives_count)
            row_sum, _ = _parse_percent_text(row_sum_item.text() if row_sum_item else "")
            if row_sum is None or row_sum < 0 or row_sum > 100:
                return f"第{row_idx-1}行小计异常,需在0-100之间"

        # Check link totals
        start_row = 2
        for link_idx, count in enumerate(self.link_counts):
            expected = 0.0 if count <= 0 else 100.0
            link_total_item = self.table.item(start_row, 3 + self.objectives_count)
            actual, _ = _parse_percent_text(link_total_item.text() if link_total_item else "")
            if actual is None or actual != expected:
                return f"{self.link_names[link_idx]}合计必须为{int(expected)}%"
            rows_for_link = 1 if count <= 0 else count
            start_row += rows_for_link

        # Check total objective weight sum
        total_row = self.table.rowCount() - 1
        total_sum_item = self.table.item(total_row, 2 + self.objectives_count)
        total_sum, _ = _parse_percent_text(total_sum_item.text() if total_sum_item else "")
        if total_sum is None or total_sum != 100.0:
            return "课程目标总权重合计必须为100%"

        return None

    def _collect_data(self):
        methods_data = []
        for row_idx, meta in self._iter_method_rows():
            method_name = self.table.item(row_idx, 1).text().strip()
            weights = []
            for obj_idx in range(self.objectives_count):
                value, _ = _parse_percent_text(self.table.item(row_idx, 2 + obj_idx).text())
                weights.append(value)
            subtotal, _ = _parse_percent_text(self.table.item(row_idx, 2 + self.objectives_count).text())
            methods_data.append(
                {
                    "link_idx": meta["link_idx"],
                    "method_name": method_name,
                    "weights": weights,
                    "subtotal": subtotal,
                }
            )

        obj_totals = []
        total_row = self.table.rowCount() - 1
        for obj_idx in range(self.objectives_count):
            value, _ = _parse_percent_text(self.table.item(total_row, 2 + obj_idx).text())
            obj_totals.append(value)
        total_sum, _ = _parse_percent_text(self.table.item(total_row, 2 + self.objectives_count).text())

        return methods_data, obj_totals, total_sum

    def _on_save(self):
        error = self._validate()
        if error:
            QMessageBox.warning(self, "提示", error)
            return
        if Document is None:
            QMessageBox.warning(self, "提示", "缺少 python-docx 依赖,无法生成 Word 文件")
            return
        methods_data, obj_totals, total_sum = self._collect_data()
        outputs_dir = _ensure_outputs_dir(os.getcwd())
        output_path = os.path.join(outputs_dir, "课程考核与课程目标对应关系表.docx")
        json_path = os.path.join(outputs_dir, "课程考核与课程目标对应关系表.json")
        try:
            export_relation_table(
                output_path,
                self.objectives_count,
                self.link_names,
                self.link_ratios,
                self.link_counts,
                methods_data,
                obj_totals,
                total_sum,
            )
            payload = export_relation_json(
                json_path,
                self.objectives_count,
                self.link_names,
                self.link_ratios,
                self.link_counts,
                methods_data,
                obj_totals,
                total_sum,
            )
        except Exception as exc:
            QMessageBox.warning(self, "提示", f"生成Word失败:{exc}")
            return
        parent = self.parent()
        if parent and hasattr(parent, "set_relation_payload"):
            parent.set_relation_payload(payload)
        QMessageBox.information(self, "完成", f"已生成:{output_path}\n{json_path}")

    def _on_back(self):
        self.close()
        parent = self.parent()
        setup = RelationTableSetupDialog(
            parent,
            default_objectives=self.objectives_count,
            default_counts=self.link_counts,
        )
        if setup.exec():
            values = setup.result_values
            if values:
                dialog = RelationTableEditorDialog(parent, *values)
                dialog.exec()


def _set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_center(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_relation_table(
    output_path: str,
    objectives_count: int,
    link_names: List[str],
    link_ratios: List[float],
    link_counts: List[int],
    methods_data: List[dict],
    obj_totals: List[float],
    total_sum: float,
):
    doc = Document()
    cols = 2 + objectives_count + 2
    method_rows = sum(1 if c <= 0 else c for c in link_counts)
    rows = 2 + method_rows + 1
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    # Header row
    table.cell(0, 0).text = "考核环节"
    table.cell(0, 1).text = "考核方式"
    table.cell(0, 2).text = "课程目标分权重"
    table.cell(0, 2).merge(table.cell(0, 1 + objectives_count))
    table.cell(0, 2 + objectives_count).text = "小计"
    table.cell(0, 3 + objectives_count).text = "合计"

    # Header row 2
    for i in range(objectives_count):
        table.cell(1, 2 + i).text = f"课程目标{i+1}"

    # Merge header vertical cells
    for col in [0, 1, 2 + objectives_count, 3 + objectives_count]:
        table.cell(0, col).merge(table.cell(1, col))

    # Fill rows
    data_row = 2
    method_idx = 0
    for link_idx, link_name in enumerate(link_names):
        rows_for_link = 1 if link_counts[link_idx] <= 0 else link_counts[link_idx]
        link_label = f"{link_name}\n{_format_percent(link_ratios[link_idx] * 100)}"
        table.cell(data_row, 0).text = link_label
        if rows_for_link > 1:
            table.cell(data_row, 0).merge(table.cell(data_row + rows_for_link - 1, 0))
            table.cell(data_row, 3 + objectives_count).merge(table.cell(data_row + rows_for_link - 1, 3 + objectives_count))
        link_total = 0.0
        for _ in range(rows_for_link):
            row = data_row
            method = methods_data[method_idx]
            table.cell(row, 1).text = method["method_name"] or " "
            for obj_idx, value in enumerate(method["weights"]):
                table.cell(row, 2 + obj_idx).text = _format_percent(value)
            table.cell(row, 2 + objectives_count).text = _format_percent(method["subtotal"])
            link_total += method["subtotal"]
            data_row += 1
            method_idx += 1
        table.cell(data_row - rows_for_link, 3 + objectives_count).text = _format_percent(link_total)

    # Total row
    total_row = rows - 1
    table.cell(total_row, 0).text = "100%"
    table.cell(total_row, 1).text = "课程目标总权重"
    for obj_idx, value in enumerate(obj_totals):
        table.cell(total_row, 2 + obj_idx).text = _format_percent(value)
    table.cell(total_row, 2 + objectives_count).text = _format_percent(total_sum)
    table.cell(total_row, 3 + objectives_count).text = _format_percent(100.0)

    # Style cells
    for row in range(rows):
        for col in range(cols):
            cell = table.cell(row, col)
            _set_paragraph_center(cell)

    doc.save(output_path)


def export_relation_json(
    output_path: str,
    objectives_count: int,
    link_names: List[str],
    link_ratios: List[float],
    link_counts: List[int],
    methods_data: List[dict],
    obj_totals: List[float],
    total_sum: float,
):
    links = []
    for link_idx, link_name in enumerate(link_names):
        methods = []
        for method in methods_data:
            if method["link_idx"] != link_idx:
                continue
            supports = {}
            for obj_idx, value in enumerate(method["weights"]):
                supports[f"课程目标{obj_idx + 1}"] = round(value / 100.0, 6)
            methods.append(
                {
                    "name": method["method_name"],
                    "supports": supports,
                    "subtotal": round(method["subtotal"] / 100.0, 6),
                }
            )
        links.append(
            {
                "name": link_name,
                "ratio": round(link_ratios[link_idx], 6),
                "methods": methods,
            }
        )

    objectives = {}
    for obj_idx, total in enumerate(obj_totals):
        objectives[f"课程目标{obj_idx + 1}"] = round(total / 100.0, 6)

    payload = {
        "objectives_count": objectives_count,
        "links": links,
        "objectives_total_weights": objectives,
        "total_sum": round(total_sum / 100.0, 6),
    }
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return payload
